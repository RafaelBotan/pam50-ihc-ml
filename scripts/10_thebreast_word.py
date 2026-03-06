"""
Final manuscript in DOCX (Word) for The Breast (Elsevier).
Generates 4 files:
  1. manuscript_TheBreast_FINAL.docx       (with inline figures & tables - for review)
  2. manuscript_TheBreast_submission.docx   (placeholders only - for Editorial Manager)
  3. title_page.docx                       (separate title page)
  4. tables.docx                           (separate tables file)

Revision notes (v2):
  - Correction #1:  Section 2.4 rewritten to explain surrogate non-classifiability in TCGA
  - Correction #2:  Sentence added to end of 2.3 (variable harmonisation)
  - Correction #3:  Table 1 title + footnotes added
  - Correction #4:  Section 3.2 rewritten (surrogate asymmetry across cohorts)
  - Correction #5:  Section 3.3 rewritten (TCGA as limited validation)
  - Correction #6:  Ethics statement added
  - Correction #7:  Table 4 retitled + explanatory note
  - Correction #8:  Data availability → choice B (upon acceptance)
  - Correction #9:  Figure 6 (information ceiling) moved to supplementary
  - Correction #10: Conclusions tightened
  - Correction #11: AI declaration made more sober
  - Correction #12: Sentence added to end of 2.7 (subset labelling)
"""

import os, sys
import pandas as pd
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_R = os.path.join(BASE_DIR, "figures_R")
FIG_PY = os.path.join(BASE_DIR, "figures")
TAB_DIR = os.path.join(BASE_DIR, "tables")
MS_DIR = os.path.join(BASE_DIR, "manuscript")
os.makedirs(MS_DIR, exist_ok=True)

t_flow = pd.read_csv(os.path.join(TAB_DIR, 'table_sample_flow.csv'))
t_lumab = pd.read_csv(os.path.join(TAB_DIR, 'table_lumAB_crossover.csv'))
t_gz = pd.read_csv(os.path.join(TAB_DIR, 'table_grey_zone.csv'))
t_sens_raw = pd.read_csv(os.path.join(TAB_DIR, 'table_sensitivity_enhanced.csv'))
# Enrich Table 4 with model/feature-set info for clarity
t_sens = t_sens_raw.copy()
t_sens.insert(1, 'Model', 'XGBoost')
t_sens.insert(2, 'Feature Set', 'Set 1 (ER, PR, HER2)')

# ── Figure mapping (correction #9: Fig 6 moved to supplementary) ──
# Main figures: 1-5 only
FIGURES = {
    1: os.path.join(FIG_R, 'fig2_cohort_characteristics.png'),
    2: os.path.join(FIG_R, 'fig4_luminal_crossover.png'),
    3: os.path.join(FIG_R, 'fig3_forest_h2h.png'),
    4: os.path.join(FIG_PY, 'fig4_confusion.png'),
    5: os.path.join(FIG_R, 'fig5_grey_zone.png'),
}

# Supplementary: S1-S5 (information ceiling now S5)
SUPP_FIGURES = {
    'S1': os.path.join(FIG_R, 'fig1_cohort_design.png'),
    'S2': os.path.join(FIG_R, 'fig6_feature_importance.png'),
    'S3': os.path.join(FIG_R, 'fig7_sensitivity_3v4.png'),
    'S4': os.path.join(FIG_R, 'fig8_feature_set_comparison.png'),
    'S5': os.path.join(FIG_R, 'fig10_information_ceiling.png'),
}

AFFILIATIONS = [
    '\u00b9 Department of Oncology, Universidade de Bras\u00edlia, Bras\u00edlia, Brazil',
    '\u00b2 Department of Proctology, Universidade de Bras\u00edlia, Bras\u00edlia, Brazil',
]

TITLE = (
    'Routine pathology cannot reliably reproduce 4-class PAM50 intrinsic subtypes '
    'in breast cancer: a multicohort in silico study defining the luminal grey zone'
)

REFS = [
    '[1] Parker JS, Mullins M, Cheang MC, et al. Supervised risk predictor of breast cancer based on intrinsic subtypes. J Clin Oncol 2009;27(8):1160\u20137.',
    '[2] Perou CM, Sorlie T, Eisen MB, et al. Molecular portraits of human breast tumours. Nature 2000;406(6797):747\u201352.',
    '[3] Prat A, Fan C, Fernandez A, et al. Clinical implications of the intrinsic molecular subtypes of breast cancer. Breast 2015;24(Suppl 2):S26\u201335.',
    '[4] Goldhirsch A, Winer EP, Coates AS, et al. Personalizing the treatment of women with early breast cancer: highlights of the St Gallen International Expert Consensus 2013. Ann Oncol 2013;24(9):2206\u201323.',
    '[5] Bastien RR, Rodriguez-Lescure A, Ebbert MT, et al. PAM50 breast cancer subtyping by RT-qPCR and concordance with standard clinical molecular markers. BMC Med Genomics 2012;5:44.',
    '[6] Prat A, Cheang MC, Martin M, et al. Prognostic significance of progesterone receptor-positive tumor cells within immunohistochemically defined luminal A breast cancer. J Clin Oncol 2013;31(2):203\u20139.',
    '[7] Collins GS, Reitsma JB, Altman DG, Moons KG. Transparent reporting of a multivariable prediction model for individual prognosis or diagnosis (TRIPOD). BMJ 2015;350:g7594.',
    '[8] Brueffer C, Vallon-Christersson J, Grabau D, et al. Clinical value of RNA sequencing-based classifiers for prediction of the five conventional breast cancer biomarkers: a SCAN-B report. JCO Precis Oncol 2018;2:1\u201318.',
    '[9] Cancer Genome Atlas Network. Comprehensive molecular portraits of human breast tumours. Nature 2012;490(7418):61\u201370.',
    '[10] Curtis C, Shah SP, Chin SF, et al. The genomic and transcriptomic architecture of 2,000 breast tumours reveals novel subgroups. Nature 2012;486(7403):346\u201352.',
    '[11] Cardoso F, van\u2019t Veer LJ, Bogaerts J, et al. 70-gene signature as an aid to treatment decisions in early-stage breast cancer. N Engl J Med 2016;375(8):717\u201329.',
    '[12] Chia SK, Bramwell VH, Tu D, et al. A 50-gene intrinsic subtype classifier for prognosis and prediction of benefit from adjuvant tamoxifen. Clin Cancer Res 2012;18(16):4465\u201372.',
]


# ── helpers ──────────────────────────────────────────────────

def set_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 2.0


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(14); r.font.name = 'Times New Roman'


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'


def para(doc, text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    return p


def bold_start(doc, bold, normal):
    p = doc.add_paragraph()
    r1 = p.add_run(bold)
    r1.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(12)
    r2 = p.add_run(normal)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)
    return p


def fig(doc, fig_path, caption, inline=True, width=Inches(5.5)):
    if inline and os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run()
        run.add_picture(fig_path, width=width)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run('[INSERT FIGURE HERE]')
        r.bold = True; r.font.size = Pt(11)
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(150, 0, 0)
    pc = doc.add_paragraph()
    pc.paragraph_format.space_after = Pt(12)
    rc = pc.add_run(caption)
    rc.font.name = 'Times New Roman'; rc.font.size = Pt(10); rc.italic = True


def tbl(doc, df, caption, inline=True, footnotes=None):
    """Add a formatted table from a DataFrame, with optional footnotes."""
    if not inline:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f'[{caption}]')
        r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(150, 0, 0)
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(caption)
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for par in cell.paragraphs:
            for run in par.runs:
                run.bold = True; run.font.size = Pt(9); run.font.name = 'Times New Roman'
    for i, (_, row) in enumerate(df.iterrows()):
        for j, col in enumerate(df.columns):
            cell = table.rows[i + 1].cells[j]
            val = str(row[col]) if pd.notna(row[col]) else ''
            cell.text = val
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9); run.font.name = 'Times New Roman'
    # Footnotes
    if footnotes:
        pn = doc.add_paragraph()
        pn.paragraph_format.space_before = Pt(4)
        rn = pn.add_run(footnotes)
        rn.font.name = 'Times New Roman'; rn.font.size = Pt(8); rn.italic = True
    doc.add_paragraph()


def add_authors(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Rafael de Negreiros Botan ')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    r2 = p.add_run('1,*')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12); r2.font.superscript = True
    r3 = p.add_run(', Jo\u00e3o Batista de Sousa ')
    r3.font.name = 'Times New Roman'; r3.font.size = Pt(12)
    r4 = p.add_run('2')
    r4.font.name = 'Times New Roman'; r4.font.size = Pt(12); r4.font.superscript = True
    doc.add_paragraph()
    for aff in AFFILIATIONS:
        pa = doc.add_paragraph()
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ra = pa.add_run(aff)
        ra.font.name = 'Times New Roman'; ra.font.size = Pt(10); ra.italic = True
    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = pc.add_run('* Corresponding author: Rafael de Negreiros Botan \u2013 oncologista@gmail.com')
    rc.font.name = 'Times New Roman'; rc.font.size = Pt(10)
    doc.add_paragraph()
    doc.add_paragraph()


# ── Table 1 footnotes (correction #3) ───────────────────────

TABLE1_FOOTNOTES = (
    'Abbreviations: PAM50, Prediction Analysis of Microarray 50; IHC, immunohistochemistry. '
    'Set 1: ER, PR, HER2. Set 2: Set 1 plus histological grade. Set 3: Set 2 plus Ki-67. '
    'Source n refers to the total cohort size available in the source dataset. '
    'N (4-class PAM50) refers to tumours with an evaluable 4-class molecular reference label '
    'after exclusion of Normal-like and claudin-low categories. '
    'N (Surrogate classifiable) refers to tumours for which full 4-class surrogate assignment '
    'was possible. In HR-positive/HER2-negative tumours, this required a proliferation variable '
    '(Ki-67 or grade) to separate Luminal A-like from Luminal B-like; therefore, surrogate '
    'applicability was reduced in TCGA-BRCA, where grade was unavailable and Ki-67 was absent. '
    'N (Head-to-head Set 1 vs Surrogate) refers to the exact matched subset classifiable by '
    'both the machine-learning model and the surrogate, used for fair denominator comparisons.'
)

# ── Table 4 footnote (correction #7) ────────────────────────

TABLE4_FOOTNOTES = (
    'All values refer to XGBoost Set 1 (ER, PR, HER2) applied to the largest model-evaluable '
    'subset in each cohort. The 4-class macro-F1 values (0.514 and 0.522) are therefore not '
    'directly comparable to the matched head-to-head results in Section 3.3, which use '
    'XGBoost Set 2 (+ grade) on a different denominator (n=1,544 in METABRIC). '
    'All direct surrogate-versus-model comparisons are reported separately using identical '
    'denominators.'
)


# ── build manuscript ─────────────────────────────────────────

def build_manuscript(inline=True):
    doc = Document()
    set_style(doc)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(TITLE)
    r.bold = True; r.font.size = Pt(14); r.font.name = 'Times New Roman'
    doc.add_paragraph()

    # Authors
    add_authors(doc)

    # ── Abstract ──
    h1(doc, 'Abstract')
    bold_start(doc, 'Background: ',
        'Immunohistochemistry-based surrogate classification is widely used to approximate '
        'intrinsic breast cancer subtypes in routine care, but its agreement with PAM50 is '
        'incomplete. We tested whether machine-learning models applied to routine '
        'clinicopathological variables could improve on the conventional surrogate and '
        'quantified the information ceiling imposed by these variables.')
    bold_start(doc, 'Methods: ',
        'Four public datasets with 4-class PAM50 labels were analysed (GSE81538, n=383; '
        'GSE96058, n=2,867; TCGA-BRCA, n=519; METABRIC, n=1,608; total n=5,377). Logistic '
        'regression, random forest, and XGBoost were trained on the SCAN-B cohorts using '
        'nested feature sets of increasing complexity: ER/PR/HER2 (Set 1), + grade (Set 2), '
        'and + Ki-67 (Set 3). External validation was performed in TCGA-BRCA and METABRIC. '
        'Head-to-head comparisons with the immunohistochemical surrogate used identical '
        'evaluable subsets, and performance was estimated with bootstrap 95% confidence intervals.')
    bold_start(doc, 'Results: ',
        'No machine-learning model outperformed the surrogate in any external comparison. '
        'On METABRIC head-to-head analysis (n=1,554), the surrogate achieved macro-F1 0.646 '
        '(95% CI 0.623\u20130.669) versus 0.559 (0.533\u20130.584) for the best model (XGBoost, Set 2). '
        'The dominant source of error was the Luminal A/Luminal B boundary: 42.5% of molecular '
        'Luminal B tumours received a Luminal A-like surrogate label. Collapsing both luminal '
        'classes increased macro-F1 to 0.767. A low-confidence grey zone comprised 13.5% of '
        'cases, was 73.6% luminal, and showed only 15.9% machine-learning accuracy.')
    bold_start(doc, 'Conclusions: ',
        'Routine clinicopathological variables impose an information ceiling that prevents '
        'reliable reproduction of 4-class PAM50 subtypes. The irreducible uncertainty '
        'concentrates at the Luminal A/Luminal B boundary, defining a clinically relevant '
        'grey zone in which molecular testing adds the greatest value.')

    # ── 1. Introduction ──
    h1(doc, '1. Introduction')
    para(doc,
        'Breast cancer is a biologically heterogeneous disease. Gene-expression profiling '
        'established the major intrinsic subtypes\u2014Luminal A, Luminal B, HER2-enriched, and '
        'Basal-like\u2014which differ in prognosis, proliferation, endocrine sensitivity, and '
        'treatment responsiveness [1,2]. Among available molecular classifiers, PAM50 remains '
        'the most widely adopted framework for intrinsic subtyping and is deeply embedded in '
        'translational research and clinical decision-making [3].')
    para(doc,
        'In routine practice, however, true molecular subtyping is often unavailable. Most '
        'centres instead rely on immunohistochemical surrogates based on ER, PR, HER2, and '
        'Ki-67, as popularised by the St Gallen consensus [4]. This approach is pragmatic, '
        'inexpensive, and globally scalable, but it is only an approximation of intrinsic '
        'biology. Discordance between immunohistochemical surrogates and PAM50 is well '
        'recognised, particularly in hormone receptor-positive disease, where the Luminal '
        'A/Luminal B distinction reflects a continuous proliferative axis that routine markers '
        'capture only imperfectly [5,6].')
    para(doc,
        'An intuitive next question is whether machine-learning methods can recover additional '
        'molecular information from the same routine variables and improve on the rule-based '
        'surrogate. If successful, a purely computational approach could refine subtype '
        'approximation without additional laboratory cost. If unsuccessful, the negative result '
        'is still clinically meaningful: it would indicate that routine pathology data contain '
        'an information ceiling that no algorithm can overcome, and that molecular testing '
        'remains irreplaceable in specific patient subgroups.')
    para(doc,
        'We therefore asked two linked questions. First, how accurately does the conventional '
        'immunohistochemical surrogate reproduce 4-class PAM50 intrinsic subtypes across '
        'multiple independent cohorts? Second, can machine-learning models built from the same '
        'routine clinicopathological variables improve on that approximation? Our aim was not '
        'to replace PAM50, but to define the practical limit of routine-data recoverability '
        'and identify the subtype boundary at which molecular testing provides the greatest '
        'incremental value.')

    # ── 2. Materials and methods ──
    h1(doc, '2. Materials and methods')

    h2(doc, '2.1 Study design')
    para(doc,
        'This was a retrospective, purely in silico, multicohort diagnostic-accuracy study. '
        'The index tests were the conventional immunohistochemical surrogate and each '
        'machine-learning classifier. The reference standard was the published 4-class PAM50 '
        'label in each cohort. No new biological samples were generated or analysed. The study '
        'was designed and reported according to TRIPOD principles for multivariable prediction '
        'model studies [7].')

    h2(doc, '2.2 Data sources and cohorts')
    para(doc, 'Four public breast cancer datasets were analysed (Table 1, Fig. 1).', indent=False)
    para(doc,
        'Development cohorts: GSE81538 (SCAN-B; source n=405, analysable 4-class subset n=383) '
        'comprised RNA-seq data with multi-observer pathology consensus for ER, PR, HER2, Ki-67, '
        'and Nottingham histological grade (NHG) [8]. GSE96058 (SCAN-B; source n=3,069, '
        'analysable 4-class subset n=2,867) comprised RNA-seq data with binary '
        'immunohistochemical status and NHG grade [8].')
    para(doc,
        'External validation cohorts: TCGA-BRCA (clinical source n=1,108, analysable 4-class '
        'subset n=519) included RNA-seq data and routine ER, PR, and HER2 assessment, with '
        'HER2-equivocal cases resolved by FISH when available; histological grade was '
        'unavailable [9]. METABRIC (clinical source n=2,509, analysable 4-class subset n=1,608) '
        'included microarray-based expression data and routine ER, PR, HER2, and NHG grade [10].')
    para(doc,
        'Normal-like and claudin-low labels were excluded to ensure a uniform 4-class endpoint '
        'across cohorts. Publicly available processed data were obtained from NCBI GEO and cBioPortal.')

    # Table 1 with footnotes (correction #3)
    tbl(doc, t_flow,
        'Table 1. Cohort characteristics and evaluable sample sizes by feature set and surrogate applicability.',
        inline, footnotes=TABLE1_FOOTNOTES)

    # ── 2.3 Variable harmonisation (correction #2: added final sentence) ──
    h2(doc, '2.3 Variable harmonisation')
    para(doc,
        'ER and PR were coded as binary positive/negative variables across all cohorts. HER2 '
        'was harmonised as binary positive/negative; equivocal cases were resolved using in situ '
        'hybridisation when available. Ki-67 was available as a continuous percentage in GSE81538 '
        'and as a binary high/low variable in GSE96058; it was unavailable in TCGA-BRCA and '
        'METABRIC. Histological grade (NHG 1\u20133) was available in GSE81538, GSE96058, and '
        'METABRIC, but not in TCGA-BRCA. Because variable granularity differed across cohorts, '
        'analyses were organised into prespecified feature sets rather than forcing a single '
        'complete-case dataset across all platforms.')
    # Correction #2: new sentence
    para(doc,
        'This heterogeneity affected not only model portability but also the feasibility of '
        'rule-based 4-class surrogate assignment, particularly in HR-positive/HER2-negative '
        'tumours requiring a proliferation variable for luminal subclassification.')

    # ── 2.4 Conventional surrogate classification (correction #1: full rewrite) ──
    h2(doc, '2.4 Conventional surrogate classification')
    para(doc,
        'The baseline comparator was a St Gallen-like immunohistochemical surrogate. Tumours '
        'were classified as Luminal A-like, Luminal B-like, HER2-positive non-luminal, or '
        'triple-negative according to hormone receptor status, HER2 status, and proliferation '
        'information.')
    para(doc,
        'Operationally, Luminal A-like was defined as HR-positive/HER2-negative with low '
        'proliferation; Luminal B-like as HR-positive/HER2-negative with high proliferation '
        'or HR-positive/HER2-positive; HER2-positive non-luminal as HR-negative/HER2-positive; '
        'and triple-negative as HR-negative/HER2-negative. Proliferation was defined as '
        'Ki-67 \u226520% when available, or grade 3 when Ki-67 was unavailable but grade was '
        'present.')
    para(doc,
        'Importantly, full 4-class surrogate assignment required a proliferation variable for '
        'HR-positive/HER2-negative tumours. Therefore, in cohorts lacking both Ki-67 and grade, '
        'HR-positive/HER2-negative cases could not be further subdivided into Luminal A-like '
        'versus Luminal B-like and were considered non-classifiable for surrogate-based 4-class '
        'assignment. This explains the lower number of surrogate-classifiable cases in TCGA-BRCA '
        'despite broader availability of Set 1 variables.')
    para(doc,
        'Cases missing ER, PR, or HER2 were also excluded from surrogate classification.')

    h2(doc, '2.5 Feature sets')
    para(doc,
        'Three nested feature sets were evaluated to reflect increasing real-world data '
        'availability: Set 1 (core IHC): ER, PR, HER2. Set 2 (IHC + grade): Set 1 plus '
        'histological grade. Set 3 (IHC + grade + Ki-67): Set 2 plus Ki-67. Set 3 was not '
        'available in external validation cohorts because Ki-67 was absent or not sufficiently '
        'granular in TCGA-BRCA and METABRIC.')

    h2(doc, '2.6 Machine-learning models')
    para(doc,
        'Three supervised multiclass classifiers were trained on the combined GSE81538 and '
        'GSE96058 development cohorts: multinomial logistic regression with balanced class '
        'weights, random forest (500 trees), and XGBoost (500 rounds, maximum depth 6, '
        'learning rate 0.1). Hyperparameters were fixed before external validation. No '
        'optimisation was performed on validation cohorts.')

    # ── 2.7 Outcomes and statistical analysis (correction #12: added final sentence) ──
    h2(doc, '2.7 Outcomes and statistical analysis')
    para(doc,
        'The primary endpoint was macro-averaged F1-score for 4-class PAM50 prediction. '
        'Secondary metrics were Cohen\u2019s kappa and balanced accuracy. All performance estimates '
        'are reported with bootstrap 95% confidence intervals derived from 1,000 stratified '
        'resamples. Head-to-head comparisons between machine-learning models and the '
        'immunohistochemical surrogate were restricted to the exact intersection of cases '
        'classifiable by both methods, ensuring identical denominators.')
    para(doc,
        'Prespecified sensitivity analyses included: (1) a 3-class formulation in which '
        'Luminal A and Luminal B were collapsed into a single luminal class; and (2) a '
        'grey-zone analysis based on the maximum predicted class probability, using low '
        '(<0.50), intermediate (0.50\u20130.70), and high (>0.70) confidence bins.')
    # Correction #12: new sentence
    para(doc,
        'Separate descriptive analyses performed on full model-evaluable subsets are presented '
        'only as supportive analyses and should not be interpreted as direct comparisons with '
        'the surrogate unless explicitly labelled as matched head-to-head.')

    # ── 3. Results ──
    h1(doc, '3. Results')

    h2(doc, '3.1 Cohort characteristics')
    para(doc,
        'After exclusions, 5,377 tumours with 4-class PAM50 labels were included: 3,250 in '
        'the development cohorts and 2,127 in external validation cohorts. Luminal A was the '
        'most prevalent subtype in every cohort, followed by Luminal B, with smaller '
        'HER2-enriched and Basal-like fractions. ER positivity ranged from 77% to 92%, and '
        'HER2 positivity from 13% to 22%. Histological grade distributions were broadly '
        'comparable across the three cohorts in which grade was available (Fig. 1). These '
        'between-cohort differences created a realistic and deliberately stringent validation '
        'setting spanning two molecular platforms and heterogeneous biomarker availability.')
    fig(doc, FIGURES[1],
        'Fig. 1. Cohort characteristics. (A) PAM50 subtype distribution by cohort. '
        '(B) ER and HER2 positivity rates. (C) Histological grade distribution. '
        '(D) Absolute counts by subtype.', inline)

    # ── 3.2 Baseline performance (correction #4: rewritten) ──
    h2(doc, '3.2 Baseline performance of the immunohistochemical surrogate')
    para(doc,
        'The immunohistochemical surrogate was classifiable in 4,694 of 5,377 tumours (87.3%), '
        'but surrogate applicability varied materially across cohorts because full 4-class '
        'assignment required a proliferation variable in HR-positive/HER2-negative disease. '
        'This limitation had little impact in METABRIC and the SCAN-B cohorts, but reduced '
        'classifiability substantially in TCGA-BRCA, where grade was unavailable and Ki-67 '
        'was absent.')
    para(doc,
        'In METABRIC head-to-head analysis (n=1,554), the surrogate achieved macro-F1 0.646 '
        '(95% CI 0.623\u20130.669) and kappa 0.471 (0.436\u20130.506). Concordance was highest for '
        'Basal-like and HER2-enriched tumours and clearly lower for luminal disease.')
    para(doc,
        'The dominant error pattern was crossover at the Luminal A/Luminal B boundary. In '
        'METABRIC, 42.5% of molecular Luminal B tumours received a Luminal A-like surrogate '
        'label, while 26.9% of molecular Luminal A tumours were overcalled as Luminal B-like '
        '(Fig. 2, Table 2). The same general pattern was observed across cohorts, although '
        'estimates in TCGA-BRCA were less stable because only a restricted subset was '
        'surrogate-classifiable.')
    fig(doc, FIGURES[2],
        'Fig. 2. Luminal A/B crossover by IHC surrogate across cohorts.', inline)
    tbl(doc, t_lumab,
        'Table 2. Luminal A/Luminal B crossover rates by cohort.', inline)

    # ── 3.3 ML did not outperform (correction #5: rewritten) ──
    h2(doc, '3.3 Machine-learning models did not outperform the surrogate')
    para(doc,
        'No machine-learning model outperformed the immunohistochemical surrogate in any '
        'external matched head-to-head comparison.', indent=False)
    para(doc,
        'In METABRIC, the best Set 1 model (XGBoost) achieved macro-F1 0.523 (95% CI '
        '0.503\u20130.544), substantially below the surrogate. Adding grade improved performance '
        'modestly: the best Set 2 model (XGBoost) reached 0.559 (0.533\u20130.584), but remained '
        'inferior to the surrogate value of 0.646 in the matched subset.')
    para(doc,
        'In TCGA-BRCA, external comparison was intrinsically constrained because only a '
        'limited subset was surrogate-classifiable in the absence of grade and Ki-67. Within '
        'this matched subset (n=162), the surrogate achieved macro-F1 0.485 (0.431\u20130.541), '
        'whereas the best machine-learning model reached 0.368. Accordingly, TCGA-BRCA should '
        'be interpreted as a structurally limited external validation cohort, while METABRIC '
        'provides the cleanest external comparison.')
    para(doc,
        'Across both external datasets, the overall pattern was consistent: grade added some '
        'recoverable signal, but this gain was insufficient to match the conventional surrogate.')
    fig(doc, FIGURES[3],
        'Fig. 3. Head-to-head macro-F1 comparison with bootstrap 95% confidence intervals.', inline)

    h2(doc, '3.4 Confusion-matrix analysis')
    para(doc,
        'Normalised confusion matrices in METABRIC (Fig. 4) showed that both approaches '
        'captured broad non-luminal biology better than fine luminal subclassification. The '
        'surrogate identified Basal-like and HER2-enriched tumours with relatively high '
        'sensitivity, but remained poor at separating Luminal A from Luminal B. XGBoost with '
        'grade modestly improved Luminal B capture at the cost of reduced Luminal A specificity, '
        'indicating a trade-off rather than a net increase in recoverable information.')
    fig(doc, FIGURES[4],
        'Fig. 4. Normalised confusion matrices on METABRIC. (A) IHC surrogate. '
        '(B) XGBoost Set 1 (ER, PR, HER2). (C) XGBoost Set 2 (+ grade).', inline)

    h2(doc, '3.5 Grey-zone analysis')
    para(doc,
        'Prediction confidence revealed a clinically meaningful low-certainty subgroup. In '
        'METABRIC, 13.5% of cases (n=208) had a maximum class probability <0.50. These tumours '
        'were predominantly luminal (73.6%) and showed only 15.9% machine-learning accuracy, '
        'while the surrogate still reached 54.8% accuracy (Fig. 5, Table 3).')
    para(doc,
        'By contrast, in the high-confidence zone (>0.70; 58.2% of cases), machine learning '
        'and the surrogate performed similarly (73.1% versus 73.8%). The grey-zone framework '
        'therefore identifies a specific subgroup\u2014predominantly HR-positive/HER2-negative '
        'disease\u2014in which routine-data-based approximation is least reliable and molecular '
        'testing is most likely to add value.')
    fig(doc, FIGURES[5],
        'Fig. 5. Grey-zone analysis on METABRIC. (A) Accuracy by confidence zone. '
        '(B) Distribution of cases. (C) Luminal A+B proportion by confidence zone.', inline)
    tbl(doc, t_gz,
        'Table 3. Grey-zone analysis by confidence tier.', inline)

    h2(doc, '3.6 Sensitivity analysis: 3-class collapse')
    para(doc,
        'Collapsing Luminal A and Luminal B into a single luminal class substantially improved '
        'performance. Using the best Set 1 model (XGBoost) on the largest evaluable subsets, '
        'macro-F1 increased from 0.514 to 0.777 in TCGA-BRCA (n=435) and from 0.522 to '
        '0.767 in METABRIC (n=1,608) (Table 4). This result is central to interpretation: '
        'most of the classification deficit resides within the luminal boundary rather than '
        'in the broader separation of luminal, HER2-enriched, and Basal-like disease.')
    para(doc,
        'The 4-class baseline values in Table 4 (0.514 and 0.522) correspond to XGBoost '
        'Set 1 applied to the full model-evaluable subsets and are therefore not directly '
        'comparable to the matched head-to-head values reported in Section 3.3 (0.559 for '
        'XGBoost Set 2, n=1,544), which use a different feature set and a different denominator.')
    # Table 4 with footnote (correction #7)
    tbl(doc, t_sens,
        'Table 4. Sensitivity analysis comparing 4-class and 3-class performance in model-evaluable subsets.',
        inline, footnotes=TABLE4_FOOTNOTES)

    # ── 3.7 Information ceiling (correction #9: figure moved to supp, text kept brief) ──
    h2(doc, '3.7 Information ceiling')
    para(doc,
        'Taken together, the results indicate an information ceiling imposed by routine '
        'clinicopathological variables. Adding grade and, where available, Ki-67 produced '
        'incremental gains, but these gains plateaued below the surrogate and well below '
        'molecular classification. The limiting factor was therefore not model choice, but '
        'information content in the inputs (Supplementary Fig. S5).')

    # ── 4. Discussion ──
    h1(doc, '4. Discussion')
    para(doc,
        'This multicohort in silico study shows that routine pathology variables do not '
        'contain enough stable information to reliably reproduce 4-class PAM50 intrinsic '
        'subtypes across independent cohorts. Three findings define this conclusion: the '
        'conventional immunohistochemical surrogate misclassified approximately one-third of '
        'tumours; most errors concentrated at the Luminal A/Luminal B boundary; and '
        'machine-learning models did not consistently improve on the surrogate in external '
        'validation.')
    para(doc,
        'The negative result is therefore the main result. Machine learning did not fail '
        'because the algorithms were intrinsically inadequate; rather, the routine variables '
        'themselves appear insufficiently informative for stable 4-class reconstruction of '
        'intrinsic biology. This distinction matters. It shifts the interpretation away from '
        'model optimisation and toward the practical limit of what routine pathology can encode.')
    para(doc,
        'The luminal boundary was the key bottleneck. Biologically, the distinction between '
        'Luminal A and Luminal B reflects a continuous proliferative gradient driven by '
        'coordinated gene-expression programs [1]. In contrast, routine pathology compresses '
        'this axis into much coarser markers such as Ki-67 and histological grade. Even when '
        'both variables were available, the gap was narrowed only modestly. The sensitivity '
        'analysis supports this interpretation: once Luminal A and Luminal B were collapsed, '
        'performance rose sharply in both validation cohorts, indicating that broad luminal '
        'versus non-luminal separation is recoverable, while fine luminal subclassification '
        'is not.')
    para(doc,
        'This finding has direct clinical implications. Tumours with clearly non-luminal '
        'phenotypes are already reasonably well captured by routine biomarkers. By contrast, '
        'hormone receptor-positive/HER2-negative tumours near the Luminal A/Luminal B boundary '
        'remain the true zone of irreducible uncertainty [11]. The grey-zone analysis helps '
        'identify where routine methods are least trustworthy and where molecular testing is '
        'most likely to alter biological interpretation.')
    para(doc,
        'Our study also offers a broader methodological lesson. Algorithmic complexity cannot '
        'fully compensate for information-poor inputs. Future efforts should focus less on '
        'increasingly complex classifiers built on the same routine variables and more on '
        'identifying what additional information\u2014continuous biomarker measures, digital '
        'pathology, or true molecular data\u2014is required to break the ceiling.')

    h2(doc, '4.1 Relation to prior work')
    para(doc,
        'Prior studies have shown incomplete concordance between immunohistochemical surrogates '
        'and PAM50, particularly in luminal disease [5,6,12]. Our results extend that literature '
        'by testing multiple machine-learning approaches rather than only rule-based surrogates, '
        'requiring strict external validation across independent cohorts and platforms, and using '
        'head-to-head comparisons on identical evaluable subsets.')

    h2(doc, '4.2 Strengths and limitations')
    para(doc,
        'Strengths include the use of four independent cohorts spanning two molecular platforms, '
        'strict external validation, matched-denominator comparisons, bootstrap confidence '
        'intervals for all core metrics, and a grey-zone framework that translates prediction '
        'uncertainty into potential clinical utility.')
    para(doc,
        'Limitations include: most routine biomarkers were available only in simplified '
        'categorical format; continuous Ki-67 was available only in the smaller SCAN-B cohort; '
        'histological grade was absent in TCGA-BRCA; PAM50 labels were obtained from public '
        'resources with cohort-specific implementations; and this was a purely in silico study '
        'without outcome-based confirmation of clinical consequences. These limitations reflect '
        'the conditions under which routine surrogate subtyping is commonly used in practice.')

    # ── 5. Conclusions (correction #10: tightened) ──
    h1(doc, '5. Conclusions')
    para(doc,
        'Routine clinicopathological variables impose an information ceiling that prevents '
        'reliable reproduction of 4-class PAM50 intrinsic subtypes. Machine-learning models '
        'do not overcome this limitation because the main constraint lies in the input data '
        'rather than the classifier. The irreducible uncertainty concentrates at the Luminal '
        'A/Luminal B boundary, defining a clinically meaningful grey zone in which molecular '
        'testing is most likely to refine biological classification and add practical value.')

    # ── Ethics statement (correction #6) ──
    h1(doc, 'Ethics statement')
    para(doc,
        'Ethics approval and informed consent were not required because this study used only '
        'publicly available, de-identified datasets and involved no direct patient contact or '
        'new biological material.', indent=False)

    # Funding
    h1(doc, 'Funding')
    para(doc,
        'This research did not receive any specific grant from funding agencies in the '
        'public, commercial, or not-for-profit sectors.', indent=False)

    # COI
    h1(doc, 'Declaration of competing interest')
    para(doc,
        'The authors declare that they have no known competing financial interests or '
        'personal relationships that could have appeared to influence the work reported '
        'in this paper.', indent=False)

    # CRediT
    h1(doc, 'CRediT authorship contribution statement')
    p = doc.add_paragraph()
    r = p.add_run('Rafael de Negreiros Botan: ')
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    r2 = p.add_run(
        'Conceptualisation, Methodology, Software, Formal analysis, Data curation, '
        'Writing \u2013 original draft, Visualisation. ')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)
    r3 = p.add_run('Jo\u00e3o Batista de Sousa: ')
    r3.bold = True; r3.font.name = 'Times New Roman'; r3.font.size = Pt(12)
    r4 = p.add_run('Supervision, Writing \u2013 review & editing.')
    r4.font.name = 'Times New Roman'; r4.font.size = Pt(12)

    # ── AI Declaration (correction #11: more sober) ──
    h1(doc, 'Declaration of generative AI and AI-assisted technologies in the manuscript preparation process')
    para(doc,
        'During the preparation of this work, the authors used generative AI tools to assist '
        'with code refinement, language editing, and structural revision of the manuscript. '
        'After using these tools, the authors reviewed and edited the content as needed and '
        'take full responsibility for the content of the published article.', indent=False)

    # ── Data availability (correction #8: choice B) ──
    h1(doc, 'Data availability')
    para(doc,
        'All data used in this study are publicly available. GSE81538 and GSE96058 are '
        'available through NCBI GEO; TCGA-BRCA and METABRIC clinical data were accessed '
        'through cBioPortal. Analysis code and reproducibility materials will be made '
        'publicly available upon acceptance.', indent=False)

    # References
    h1(doc, 'References')
    for ref in REFS:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10)

    # ── Supplementary Material (correction #9: now includes Fig S5 = information ceiling) ──
    doc.add_page_break()
    h1(doc, 'Supplementary Material')
    supp_captions = {
        'S1': 'Supplementary Fig. S1. Study design and sample flow across cohorts.',
        'S2': 'Supplementary Fig. S2. Feature importance (Random Forest and XGBoost).',
        'S3': 'Supplementary Fig. S3. Sensitivity analysis: 4-class versus 3-class macro-F1.',
        'S4': 'Supplementary Fig. S4. Feature set comparison across cohorts.',
        'S5': ('Supplementary Fig. S5. Conceptual summary of the information ceiling. '
               'Performance increases modestly with additional routine features but remains '
               'below the molecular reference, supporting the interpretation that the principal '
               'limitation lies in input information content rather than model architecture.'),
    }
    for key in ['S1', 'S2', 'S3', 'S4', 'S5']:
        fig(doc, SUPP_FIGURES[key], supp_captions[key], inline)

    return doc


# ── title page ───────────────────────────────────────────────

def build_title_page():
    tp = Document()
    set_style(tp)

    p = tp.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('TITLE PAGE')
    r.bold = True; r.font.size = Pt(14); r.font.name = 'Times New Roman'
    tp.add_paragraph()

    p = tp.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True; r.font.size = Pt(14); r.font.name = 'Times New Roman'
    tp.add_paragraph()

    add_authors(tp)

    p = tp.add_paragraph()
    r = p.add_run('Running title: ')
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    r2 = p.add_run('Information ceiling for routine-data approximation of PAM50')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11); r2.italic = True
    tp.add_paragraph()

    p = tp.add_paragraph()
    r = p.add_run('Keywords: ')
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    r2 = p.add_run('breast cancer; PAM50; intrinsic subtype; immunohistochemistry; machine learning; grey zone')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11)
    tp.add_paragraph()

    p = tp.add_paragraph()
    r = p.add_run('Word count: ')
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    r2 = p.add_run('Abstract: ~250 words; Main text: ~2,200 words')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11)

    return tp


# ── tables file ──────────────────────────────────────────────

def build_tables():
    td = Document()
    set_style(td)
    h1(td, 'Tables')
    tbl(td, t_flow,
        'Table 1. Cohort characteristics and evaluable sample sizes by feature set and surrogate applicability.',
        True, footnotes=TABLE1_FOOTNOTES)
    tbl(td, t_lumab,
        'Table 2. Luminal A/Luminal B crossover rates by cohort.', True)
    tbl(td, t_gz,
        'Table 3. Grey-zone analysis by confidence tier.', True)
    tbl(td, t_sens,
        'Table 4. Sensitivity analysis comparing 4-class and 3-class performance in model-evaluable subsets.',
        True, footnotes=TABLE4_FOOTNOTES)
    return td


# ── main ─────────────────────────────────────────────────────

if __name__ == '__main__':
    # Title page
    tp = build_title_page()
    tp_path = os.path.join(MS_DIR, 'title_page.docx')
    tp.save(tp_path)
    print(f"Title page: {tp_path}")

    # Tables
    td = build_tables()
    td_path = os.path.join(MS_DIR, 'tables.docx')
    td.save(td_path)
    print(f"Tables: {td_path}")

    # Manuscript WITH figures and tables inline (for review)
    doc_full = build_manuscript(inline=True)
    full_path = os.path.join(MS_DIR, 'manuscript_TheBreast_FINAL.docx')
    doc_full.save(full_path)
    print(f"Manuscript (with figures): {full_path}  ({os.path.getsize(full_path)/1024:.0f} KB)")

    # Manuscript WITHOUT figures/tables inline (for Editorial Manager submission)
    doc_sub = build_manuscript(inline=False)
    sub_path = os.path.join(MS_DIR, 'manuscript_TheBreast_submission.docx')
    doc_sub.save(sub_path)
    print(f"Manuscript (submission):   {sub_path}  ({os.path.getsize(sub_path)/1024:.0f} KB)")

    print(f"\nAll files saved in {MS_DIR}/")
    print(f"  1. manuscript_TheBreast_FINAL.docx       (with inline figures & tables)")
    print(f"  2. manuscript_TheBreast_submission.docx   (placeholders - for Editorial Manager)")
    print(f"  3. title_page.docx                       (separate title page)")
    print(f"  4. tables.docx                           (separate tables)")
    print(f"  5. figures_R/*.tiff                       (600 dpi figures for separate upload)")
    print(f"\nRevision v2 corrections applied:")
    print(f"  #1  Section 2.4 rewritten (surrogate + TCGA explanation)")
    print(f"  #2  Sentence added to 2.3 (harmonisation → classifiability)")
    print(f"  #3  Table 1 footnotes added")
    print(f"  #4  Section 3.2 rewritten (surrogate asymmetry)")
    print(f"  #5  Section 3.3 rewritten (TCGA as limited validation)")
    print(f"  #6  Ethics statement added")
    print(f"  #7  Table 4 retitled + footnote")
    print(f"  #8  Data availability → upon acceptance")
    print(f"  #9  Figure 6 moved to Supplementary Fig. S5")
    print(f"  #10 Conclusions tightened")
    print(f"  #11 AI declaration made sober")
    print(f"  #12 Subset-labelling sentence added to 2.7")
