"""Build a Breast Cancer Research and Treatment review DOCX from manuscript_BCRT_v1.md."""

from pathlib import Path
import csv
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image


BASE_DIR = Path(__file__).resolve().parents[1]
MANUSCRIPT_DIR = BASE_DIR / "manuscript"
TABLE_DIR = BASE_DIR / "tables"
FIG_R = BASE_DIR / "figures_R"
FIG_PY = BASE_DIR / "figures"

SRC_MD = MANUSCRIPT_DIR / "manuscript_BCRT_v1.md"
OUT_DOCX = MANUSCRIPT_DIR / "manuscript_BCRT_v1.docx"
OUT_TABLES = MANUSCRIPT_DIR / "tables_BCRT.docx"
FIG4_TIFF = FIG_R / "fig4_confusion.tiff"


TABLES = {
    1: TABLE_DIR / "table_sample_flow.csv",
    2: TABLE_DIR / "table_lumAB_crossover.csv",
    3: TABLE_DIR / "table_grey_zone.csv",
    4: TABLE_DIR / "table_sensitivity_enhanced.csv",
}

FIGURES = {
    "1": FIG_R / "fig2_cohort_characteristics.png",
    "2": FIG_R / "fig4_luminal_crossover.png",
    "3": FIG_R / "fig3_forest_h2h.png",
    "4": FIG_PY / "fig4_confusion.png",
    "5": FIG_R / "fig5_grey_zone.png",
    "S1": FIG_R / "fig1_cohort_design.png",
    "S2": FIG_R / "fig6_feature_importance.png",
    "S3": FIG_R / "fig7_sensitivity_3v4.png",
    "S4": FIG_R / "fig8_feature_set_comparison.png",
    "S5": FIG_R / "fig10_information_ceiling.png",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=10, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_markdown_runs(paragraph, text, size=10):
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=size, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=size, italic=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size)


def add_para(doc, text="", style=None, size=10, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    add_markdown_runs(p, text, size=size)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=14 if level == 1 else 12 if level == 2 else 10, bold=True)
    return p


def read_csv_table(path):
    with path.open(newline="", encoding="utf-8-sig") as f:
        return list(csv.reader(f))


def add_table_from_csv(doc, path):
    rows = read_csv_table(path)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_font(run, size=8, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, "D9EAF7")
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph()
    return table


def add_figure(doc, fig_key, caption):
    path = FIGURES[fig_key]
    if not path.exists():
        add_para(doc, f"[MISSING FIGURE FILE: {path}]", size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width = Inches(6.3 if fig_key != "4" else 6.6)
    p.add_run().add_picture(str(path), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    add_markdown_runs(cap, caption, size=9)


def add_table_block(doc, table_no, caption):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(4)
    add_markdown_runs(cap, caption, size=9)
    add_table_from_csv(doc, TABLES[table_no])


def apply_document_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(10)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run_el = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = "1"
    run_el.append(text_el)
    fld.append(run_el)
    p._p.append(fld)


def convert_fig4_tiff():
    src = FIG_PY / "fig4_confusion.png"
    if FIG4_TIFF.exists() and src.exists() and FIG4_TIFF.stat().st_mtime >= src.stat().st_mtime:
        return
    if src.exists():
        im = Image.open(src).convert("RGB")
        im.save(FIG4_TIFF, dpi=(300, 300), compression="tiff_lzw")


def build_docx():
    convert_fig4_tiff()
    doc = Document()
    apply_document_defaults(doc)

    lines = SRC_MD.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        i += 1
        if not line or line == "---":
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            add_markdown_runs(p, line[2:], size=14)
            for run in p.runs:
                run.bold = True
            continue

        h = re.match(r"^(#{2,4})\s+(.+)$", line)
        if h:
            add_heading(doc, h.group(2), len(h.group(1)) - 1)
            continue

        table_match = re.match(r"^\*\*Table (\d+)\.\*\*\s*(.+)$", line)
        if table_match:
            table_no = int(table_match.group(1))
            add_table_block(doc, table_no, line)
            continue

        fig_match = re.match(r"^\*\*Fig\. ([1-5])\.\*\*\s*(.+)$", line)
        if fig_match:
            add_figure(doc, fig_match.group(1), line)
            continue

        supp_match = re.match(r"^\*\*Supplementary Fig\. (S[1-5])\.\*\*\s*(.+)$", line)
        if supp_match:
            add_figure(doc, supp_match.group(1), line)
            continue

        add_para(doc, line, size=10)

    doc.save(OUT_DOCX)
    build_tables_docx()


def build_tables_docx():
    doc = Document()
    apply_document_defaults(doc)
    add_heading(doc, "Tables", 1)
    captions = {
        1: "Table 1. Cohort characteristics and evaluable sample sizes by pathology feature set and surrogate applicability.",
        2: "Table 2. Luminal A/Luminal B crossover rates by cohort.",
        3: "Table 3. Grey-zone analysis by confidence tier.",
        4: "Table 4. Sensitivity analysis comparing 4-class XGBoost receptor-only and 3-class luminal-grouped performance in model-evaluable subsets.",
    }
    for table_no in [1, 2, 3, 4]:
        add_para(doc, captions[table_no], size=9)
        add_table_from_csv(doc, TABLES[table_no])
    doc.save(OUT_TABLES)


if __name__ == "__main__":
    build_docx()
    print(f"Manuscript DOCX: {OUT_DOCX}")
    print(f"Tables DOCX: {OUT_TABLES}")
    print(f"Generated Fig4 TIFF: {FIG4_TIFF.exists()} {FIG4_TIFF}")
